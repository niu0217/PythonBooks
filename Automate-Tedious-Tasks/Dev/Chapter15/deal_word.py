import docx
import docx.shared
import docx.enum.text


def readWord():
    """简单的读取word文档"""
    doc = docx.Document('sources/demo.docx')
    print(f'paragraphs: {len(doc.paragraphs)}\n\n')

    # 打印所有的paragraphs
    print('paragraphs begin...')
    for i in range(len(doc.paragraphs)):
        paragraphText = doc.paragraphs[i].text
        print(paragraphText)
    print('paragraphs end...\n\n')

    # 打印paragraphs[1]的所有runs
    print('paragraphs[1] runs begin...')
    for i in range(len(doc.paragraphs[1].runs)):
        runText = doc.paragraphs[1].runs[i].text
        print(runText)
    print('paragraphs[1] runs end...\n\n')


def getText(filename):
    """获取word中的所有文本"""
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)

    return '\n'.join(fullText)


def writeWordFirstMethod():
    """写word的第一种方式"""
    doc = docx.Document()
    doc.add_paragraph('Hello niu0217')
    doc.save('hello.docx')


def writeWordSecondMethod():
    """写word的第二种方式"""
    doc = docx.Document()
    doc.add_paragraph('Hello niu0217')
    paraObj1 = doc.add_paragraph('This is a yet second paragraph.')
    paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
    paraObj1.add_run('This text is being added to the second paragraph.')
    doc.save('multipleParagraph.docx')


def addTitle():
    """添加标题"""
    doc = docx.Document()
    doc.addheading('Header 0', 0)
    doc.addheading('Header 1', 1)
    doc.addheading('Header 2', 2)
    doc.addheading('Header 3', 3)
    doc.addheading('Header 4', 4)
    doc.save('headings.docx')


def addNewline():
    """换行"""
    doc = docx.Document()
    doc.add_paragraph('Hello niu0217')
    doc.paragraphs[0].runs[0].add_break()
    doc.add_paragraph('Hello niu0217')
    doc.save('newline.docx')


def addNewpage():
    """换页"""
    doc = docx.Document()
    doc.add_paragraph('Hello niu0217')
    doc.paragraphs[0].runs[0].add_break(
        docx.enum.text.WD_BREAK.PAGE
    )
    doc.add_paragraph('Hello niu0217')
    doc.save('newpage.docx')


def addPicture():
    """添加图片"""
    doc = docx.Document()
    doc.add_paragraph('Hello niu0217')
    doc.add_picture(
        'sources/zophie.png',
        width=docx.shared.Inches(1),
        height=docx.shared.Cm(4)
    )
    doc.save('newpicture.docx')


if __name__ == '__main__':
    # readWord()

    # result = getText('sources/demo.docx')
    # print(result)

    # writeWordFirstMethod()

    # writeWordSecondMethod()

    addTitle()
    addNewpage()
    addPicture()
